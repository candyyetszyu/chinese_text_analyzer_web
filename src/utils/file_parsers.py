# -*- coding: utf-8 -*-
"""
Extended File Format Parser Module
支持PDF、Word、網頁等格式的文件解析
"""

import os
import sys
import requests
from pathlib import Path

# PDF parsing
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Word document parsing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Web scraping
try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

# File type detection
try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False

try:
    import mimetypes
    MIMETYPES_AVAILABLE = True
except ImportError:
    MIMETYPES_AVAILABLE = False

# Add project root to path
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

class ExtendedFileParser:
    """擴展文件格式解析器"""
    
    def __init__(self):
        self.supported_formats = {
            'txt': self.parse_text_file,
            'pdf': self.parse_pdf,
            'docx': self.parse_docx,
            'doc': self.parse_docx,
            'html': self.parse_html,
            'htm': self.parse_html,
            'url': self.parse_url,
            'md': self.parse_markdown,
            'csv': self.parse_csv_as_text,
            'json': self.parse_json_as_text
        }
        
        self.capabilities = {
            'pdf': PDF_AVAILABLE,
            'docx': DOCX_AVAILABLE,
            'web_scraping': BS4_AVAILABLE,
            'file_detection': MAGIC_AVAILABLE or MIMETYPES_AVAILABLE
        }
        
        print("文件解析器初始化完成")
        print("支持的功能:")
        for feature, available in self.capabilities.items():
            status = "✓" if available else "✗"
            print(f"  {status} {feature}")
    
    def detect_file_type(self, file_path):
        """檢測文件類型"""
        if MAGIC_AVAILABLE:
            try:
                mime_type = magic.from_file(file_path, mime=True)
                return self._mime_to_extension(mime_type)
            except Exception as e:
                print(f"魔數檢測失敗: {e}")
        
        if MIMETYPES_AVAILABLE:
            mime_type, _ = mimetypes.guess_type(file_path)
            if mime_type:
                return self._mime_to_extension(mime_type)
        
        # 回退到文件擴展名
        return Path(file_path).suffix.lower().lstrip('.')
    
    def _mime_to_extension(self, mime_type):
        """MIME類型轉文件擴展名"""
        mime_map = {
            'text/plain': 'txt',
            'application/pdf': 'pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'application/msword': 'doc',
            'text/html': 'html',
            'text/markdown': 'md',
            'text/csv': 'csv',
            'application/json': 'json'
        }
        return mime_map.get(mime_type, 'unknown')
    
    def parse_file(self, file_path):
        """根據文件類型自動解析文件"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 檢測文件類型
        file_type = self.detect_file_type(file_path)
        
        # 獲取對應的解析函數
        parser_func = self.supported_formats.get(file_type)
        
        if parser_func is None:
            # 嘗試作為文本文件處理
            try:
                return self.parse_text_file(file_path)
            except Exception as e:
                raise ValueError(f"不支持的文件格式: {file_type}")
        
        try:
            return parser_func(file_path)
        except Exception as e:
            print(f"解析文件 {file_path} 時出錯: {e}")
            # 嘗試作為文本文件處理
            try:
                return self.parse_text_file(file_path)
            except Exception as text_error:
                raise ValueError(f"無法解析文件: {e}")
    
    def parse_text_file(self, file_path):
        """解析純文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    content = f.read()
            except UnicodeDecodeError:
                with open(file_path, 'r', encoding='latin-1') as f:
                    content = f.read()
        
        return {
            'content': content,
            'metadata': {
                'file_type': 'text',
                'file_path': file_path,
                'encoding': 'auto-detected'
            }
        }
    
    def parse_pdf(self, file_path):
        """解析PDF文件"""
        if not PDF_AVAILABLE:
            raise ImportError("PDF解析庫不可用，請安裝 PyPDF2 和 pdfplumber")
        
        content = ""
        metadata = {
            'file_type': 'pdf',
            'file_path': file_path,
            'total_pages': 0,
            'parser': 'unknown'
        }
        
        # 嘗試使用pdfplumber（更好的表格和佈局支持）
        try:
            with pdfplumber.open(file_path) as pdf:
                metadata['total_pages'] = len(pdf.pages)
                metadata['parser'] = 'pdfplumber'
                
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        content += page_text + "\n"
                        
        except Exception as e:
            print(f"pdfplumber解析失敗，嘗試PyPDF2: {e}")
            
            # 回退到PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata['total_pages'] = len(pdf_reader.pages)
                    metadata['parser'] = 'PyPDF2'
                    
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            content += page_text + "\n"
                            
            except Exception as pdf2_error:
                raise ValueError(f"PDF解析失敗: {pdf2_error}")
        
        return {
            'content': content.strip(),
            'metadata': metadata
        }
    
    def parse_docx(self, file_path):
        """解析Word文檔"""
        if not DOCX_AVAILABLE:
            raise ImportError("Word文檔解析庫不可用，請安裝 python-docx")
        
        try:
            doc = Document(file_path)
            content = ""
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content += " | ".join(row_text) + "\n"
            
            metadata = {
                'file_type': 'docx',
                'file_path': file_path,
                'paragraphs_count': len(doc.paragraphs),
                'tables_count': len(doc.tables)
            }
            
            return {
                'content': content.strip(),
                'metadata': metadata
            }
            
        except Exception as e:
            raise ValueError(f"Word文檔解析失敗: {e}")
    
    def parse_html(self, file_path):
        """解析HTML文件"""
        if not BS4_AVAILABLE:
            raise ImportError("HTML解析庫不可用，請安裝 beautifulsoup4")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    html_content = f.read()
            except UnicodeDecodeError:
                with open(file_path, 'r', encoding='latin-1') as f:
                    html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 移除script和style標籤
        for script in soup(["script", "style"]):
            script.decompose()
        
        # 提取文本
        text = soup.get_text()
        
        # 清理文本
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        content = '\n'.join(chunk for chunk in chunks if chunk)
        
        metadata = {
            'file_type': 'html',
            'file_path': file_path,
            'title': soup.title.string if soup.title else None
        }
        
        return {
            'content': content,
            'metadata': metadata
        }
    
    def parse_url(self, url):
        """解析網頁URL"""
        if not BS4_AVAILABLE:
            raise ImportError("網頁解析庫不可用，請安裝 beautifulsoup4")
        
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            
            response = requests.get(url, headers=headers, timeout=30)
            response.encoding = response.apparent_encoding
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # 移除script和style標籤
            for script in soup(["script", "style"]):
                script.decompose()
            
            # 提取文本
            text = soup.get_text()
            
            # 清理文本
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            content = '\n'.join(chunk for chunk in chunks if chunk)
            
            metadata = {
                'file_type': 'web_page',
                'url': url,
                'title': soup.title.string if soup.title else None,
                'status_code': response.status_code,
                'content_type': response.headers.get('content-type', 'unknown')
            }
            
            return {
                'content': content,
                'metadata': metadata
            }
            
        except requests.RequestException as e:
            raise ValueError(f"網頁獲取失敗: {e}")
        except Exception as e:
            raise ValueError(f"網頁解析失敗: {e}")
    
    def parse_markdown(self, file_path):
        """解析Markdown文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    content = f.read()
            except UnicodeDecodeError:
                with open(file_path, 'r', encoding='latin-1') as f:
                    content = f.read()
        
        # 簡單的Markdown標記清理
        import re
        
        # 移除標題標記
        content = re.sub(r'^#+\s*', '', content, flags=re.MULTILINE)
        
        # 移除鏈接語法，保留文本
        content = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', content)
        
        # 移除圖片語法
        content = re.sub(r'!\[([^\]]*)\]\([^\)]+\)', r'\1', content)
        
        # 移除粗體和斜體標記
        content = re.sub(r'\*\*([^\*]+)\*\*', r'\1', content)
        content = re.sub(r'\*([^\*]+)\*', r'\1', content)
        
        # 移除代碼塊
        content = re.sub(r'```[\s\S]*?```', '', content)
        content = re.sub(r'`([^`]+)`', r'\1', content)
        
        metadata = {
            'file_type': 'markdown',
            'file_path': file_path
        }
        
        return {
            'content': content.strip(),
            'metadata': metadata
        }
    
    def parse_csv_as_text(self, file_path):
        """解析CSV文件為文本"""
        try:
            import pandas as pd
            df = pd.read_csv(file_path)
            content = df.to_string(index=False)
        except ImportError:
            # 如果沒有pandas，簡單處理CSV
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except Exception as e:
            # 回退到文本解析
            return self.parse_text_file(file_path)
        
        metadata = {
            'file_type': 'csv',
            'file_path': file_path
        }
        
        return {
            'content': content,
            'metadata': metadata
        }
    
    def parse_json_as_text(self, file_path):
        """解析JSON文件為文本"""
        import json
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # 將JSON轉換為可讀文本
            content = json.dumps(data, ensure_ascii=False, indent=2)
            
        except Exception as e:
            # 回退到文本解析
            return self.parse_text_file(file_path)
        
        metadata = {
            'file_type': 'json',
            'file_path': file_path
        }
        
        return {
            'content': content,
            'metadata': metadata
        }
    
    def batch_parse_files(self, file_paths):
        """批量解析多個文件"""
        results = {}
        
        for file_path in file_paths:
            try:
                result = self.parse_file(file_path)
                results[file_path] = result
            except Exception as e:
                results[file_path] = {
                    'content': '',
                    'metadata': {
                        'error': str(e),
                        'file_path': file_path
                    }
                }
                print(f"解析文件 {file_path} 失敗: {e}")
        
        return results
    
    def get_supported_extensions(self):
        """獲取支持的文件擴展名列表"""
        return list(self.supported_formats.keys())
    
    def is_supported(self, file_path):
        """檢查文件是否支持"""
        extension = Path(file_path).suffix.lower().lstrip('.')
        return extension in self.supported_formats 